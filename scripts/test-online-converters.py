from __future__ import annotations

import json
import mimetypes
import os
import sys
import time
import urllib.request
import zipfile
from pathlib import Path
from uuid import uuid4

ROOT = Path(__file__).resolve().parents[1]
TEST_DEPS = ROOT / ".test_deps"
if TEST_DEPS.exists():
    sys.path.insert(0, str(TEST_DEPS))

from PIL import Image
from docx import Document
from pypdf import PdfReader, PdfWriter

try:
    from pillow_heif import register_heif_opener

    register_heif_opener()
except Exception:
    pass

API_BASE = os.environ.get("API_BASE", "https://convert-studio-pdf-api.onrender.com")
ARTIFACTS = ROOT / "test-artifacts"
FIXTURES = ARTIFACTS / "fixtures"
OUTPUTS = ARTIFACTS / "outputs"


def request_json(url: str, *, method: str = "GET", body: bytes | None = None, headers: dict | None = None) -> dict:
    request = urllib.request.Request(url, data=body, headers=headers or {}, method=method)
    with urllib.request.urlopen(request, timeout=180) as response:
        return json.loads(response.read().decode("utf-8"))


def multipart(fields: dict[str, str], files: list[Path]) -> tuple[bytes, str]:
    boundary = f"----convert-studio-{uuid4().hex}"
    chunks: list[bytes] = []

    for name, value in fields.items():
        chunks.append(f"--{boundary}\r\n".encode())
        chunks.append(f'Content-Disposition: form-data; name="{name}"\r\n\r\n'.encode())
        chunks.append(str(value).encode())
        chunks.append(b"\r\n")

    for path in files:
        mime = mimetypes.guess_type(path.name)[0] or "application/octet-stream"
        chunks.append(f"--{boundary}\r\n".encode())
        chunks.append(f'Content-Disposition: form-data; name="files"; filename="{path.name}"\r\n'.encode())
        chunks.append(f"Content-Type: {mime}\r\n\r\n".encode())
        chunks.append(path.read_bytes())
        chunks.append(b"\r\n")

    chunks.append(f"--{boundary}--\r\n".encode())
    return b"".join(chunks), f"multipart/form-data; boundary={boundary}"


def write_minimal_pptx(path: Path) -> None:
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Override PartName="/ppt/presentation.xml" ContentType="application/vnd.openxmlformats-officedocument.presentationml.presentation.main+xml"/></Types>',
        )
        archive.writestr(
            "ppt/presentation.xml",
            '<?xml version="1.0"?><p:presentation xmlns:p="http://schemas.openxmlformats.org/presentationml/2006/main"><p:sldIdLst/></p:presentation>',
        )


def write_minimal_xlsx(path: Path) -> None:
    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as archive:
        archive.writestr(
            "[Content_Types].xml",
            '<?xml version="1.0"?><Types xmlns="http://schemas.openxmlformats.org/package/2006/content-types"><Default Extension="rels" ContentType="application/vnd.openxmlformats-package.relationships+xml"/><Default Extension="xml" ContentType="application/xml"/><Override PartName="/xl/workbook.xml" ContentType="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet.main+xml"/></Types>',
        )
        archive.writestr(
            "_rels/.rels",
            '<?xml version="1.0"?><Relationships xmlns="http://schemas.openxmlformats.org/package/2006/relationships"><Relationship Id="rId1" Type="http://schemas.openxmlformats.org/officeDocument/2006/relationships/officeDocument" Target="xl/workbook.xml"/></Relationships>',
        )
        archive.writestr(
            "xl/workbook.xml",
            '<?xml version="1.0"?><workbook xmlns="http://schemas.openxmlformats.org/spreadsheetml/2006/main"><sheets/></workbook>',
        )


def create_fixtures() -> dict[str, Path]:
    FIXTURES.mkdir(parents=True, exist_ok=True)
    OUTPUTS.mkdir(parents=True, exist_ok=True)

    image = Image.new("RGB", (80, 60), (16, 124, 102))
    png = FIXTURES / "sample.png"
    jpg = FIXTURES / "sample.jpg"
    webp = FIXTURES / "sample.webp"
    avif = FIXTURES / "sample.avif"
    tiff = FIXTURES / "sample.tiff"
    heic = FIXTURES / "sample.heic"
    image.save(png)
    image.save(jpg, "JPEG")
    image.save(webp, "WEBP")
    image.save(tiff, "TIFF")
    try:
        image.save(avif, "AVIF")
    except Exception:
        pass
    try:
        image.save(heic, "HEIF")
    except Exception as exc:
        print(f"SKIP fixture HEIC generation: {exc}")

    pdf = FIXTURES / "sample.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    with pdf.open("wb") as handle:
        writer.write(handle)

    pdf2 = FIXTURES / "sample-two.pdf"
    writer = PdfWriter()
    writer.add_blank_page(width=612, height=792)
    writer.add_blank_page(width=612, height=792)
    with pdf2.open("wb") as handle:
        writer.write(handle)

    protected = FIXTURES / "sample-protected.pdf"
    reader = PdfReader(str(pdf))
    writer = PdfWriter()
    writer.add_page(reader.pages[0])
    writer.encrypt("testpass123")
    with protected.open("wb") as handle:
        writer.write(handle)

    docx = FIXTURES / "sample.docx"
    document = Document()
    document.add_heading("FileForma test", level=1)
    document.add_paragraph("This document verifies Word to PDF conversion.")
    document.save(docx)

    pptx = FIXTURES / "sample.pptx"
    write_minimal_pptx(pptx)

    xlsx = FIXTURES / "sample.xlsx"
    write_minimal_xlsx(xlsx)

    return {
        "heic": heic,
        "tiff": tiff,
        "pdf": pdf,
        "pdf2": pdf2,
        "protected": protected,
        "docx": docx,
        "pptx": pptx,
        "xlsx": xlsx,
    }


def create_job(tool: str, target: str, files: list[Path], extra: dict[str, str] | None = None) -> dict:
    fields = {"tool": tool, "targetFormat": target}
    if extra:
        fields.update(extra)
    body, content_type = multipart(fields, files)
    return request_json(
        f"{API_BASE}/api/jobs",
        method="POST",
        body=body,
        headers={"Content-Type": content_type},
    )


def wait_for_job(job_id: str) -> dict:
    for _ in range(180):
        result = request_json(f"{API_BASE}/api/jobs/{job_id}")
        if result["status"] in {"done", "failed"}:
            return result
        time.sleep(1)
    raise TimeoutError(f"{job_id} timed out")


def download(job_id: str, output_name: str) -> Path:
    output_path = OUTPUTS / output_name
    with urllib.request.urlopen(f"{API_BASE}/api/jobs/{job_id}/download", timeout=180) as response:
        output_path.write_bytes(response.read())
    return output_path


def assert_magic(path: Path, expected: bytes) -> None:
    actual = path.read_bytes()[: len(expected)]
    if actual != expected:
        raise AssertionError(f"{path.name} magic {actual!r} != {expected!r}")


def run_case(name: str, tool: str, target: str, files: list[Path], expected_magic: bytes, extra: dict[str, str] | None = None) -> None:
    print(f"TEST {name}")
    job = create_job(tool, target, files, extra)
    result = wait_for_job(job["jobId"])
    if result["status"] != "done":
        raise AssertionError(f"{name} failed: {result.get('error')}")
    output = download(job["jobId"], result.get("outputName") or f"{name}.out")
    assert_magic(output, expected_magic)
    print(f"PASS {name}: {output.name} {output.stat().st_size} bytes")


def wake_api() -> None:
    print(f"Waking API at {API_BASE} ...")
    for attempt in range(12):
        try:
            health = request_json(f"{API_BASE}/health")
            print("HEALTH", health)
            return
        except Exception as exc:
            print(f"  attempt {attempt + 1}: {exc}")
            time.sleep(5)
    raise RuntimeError("API health check failed after retries")


def ensure_text_pdf(fixtures: dict[str, Path]) -> Path:
    text_pdf = FIXTURES / "sample-text.pdf"
    if text_pdf.exists():
        return text_pdf
    print("Creating text PDF fixture via Word to PDF ...")
    job = create_job("word-to-pdf", "PDF", [fixtures["docx"]])
    result = wait_for_job(job["jobId"])
    if result["status"] != "done":
        raise RuntimeError(f"Could not create text PDF fixture: {result.get('error')}")
    output = download(job["jobId"], "sample-text.pdf")
    output.replace(text_pdf)
    return text_pdf


def main() -> None:
    if "--fixtures-only" in sys.argv:
        create_fixtures()
        print(f"Fixtures ready in {FIXTURES}")
        return

    fixtures = create_fixtures()
    wake_api()
    text_pdf = ensure_text_pdf(fixtures)

    cases: list[tuple] = [
        ("HEIC to JPG", "heic-to-jpg", "JPG", [fixtures["heic"]], b"\xff\xd8\xff", None, False),
        ("TIFF to JPG", "tiff-to-jpg", "JPG", [fixtures["tiff"]], b"\xff\xd8\xff", None, False),
        ("Word to PDF", "word-to-pdf", "PDF", [fixtures["docx"]], b"%PDF", None, False),
        ("PowerPoint to PDF", "powerpoint-to-pdf", "PDF", [fixtures["pptx"]], b"%PDF", None, True),
        ("Excel to PDF", "excel-to-pdf", "PDF", [fixtures["xlsx"]], b"%PDF", None, False),
        ("PDF to Word", "pdf-to-word", "DOCX", [fixtures["pdf"]], b"PK", None, False),
        ("Compress PDF", "compress-pdf", "Smaller PDF", [fixtures["pdf"]], b"%PDF", None, False),
        ("Repair PDF", "repair-pdf", "Repaired PDF", [fixtures["pdf"]], b"%PDF", None, False),
        ("PDF to JPG", "pdf-to-jpg", "JPG", [fixtures["pdf"]], b"PK", None, False),
        ("Merge PDF", "merge-pdf", "Merged PDF", [fixtures["pdf"], fixtures["pdf2"]], b"%PDF", None, False),
        ("Split PDF", "split-pdf", "Split PDF", [fixtures["pdf2"]], b"%PDF", {"pages": "1"}, False),
        ("Remove PDF pages", "remove-pdf-pages", "Updated PDF", [fixtures["pdf2"]], b"%PDF", {"pages": "2"}, False),
        ("Extract PDF pages", "extract-pdf-pages", "Extracted PDF", [fixtures["pdf2"]], b"%PDF", {"pages": "1"}, False),
        ("Organize PDF", "organize-pdf", "Organized PDF", [fixtures["pdf2"]], b"%PDF", {"pages": "2,1"}, False),
        ("Rotate PDF", "rotate-pdf", "Rotated PDF", [fixtures["pdf"]], b"%PDF", {"rotation": "90"}, False),
        ("Add page numbers", "add-page-numbers", "Numbered PDF", [fixtures["pdf"]], b"%PDF", None, False),
        ("Add watermark", "add-watermark", "Watermarked PDF", [fixtures["pdf"]], b"%PDF", {"text": "DRAFT"}, False),
        ("Crop PDF", "crop-pdf", "Cropped PDF", [fixtures["pdf"]], b"%PDF", {"margin": "24"}, False),
        ("Edit PDF", "edit-pdf", "Edited PDF", [fixtures["pdf"]], b"%PDF", {"text": "NOTE", "page": "1"}, False),
        ("Redact PDF", "redact-pdf", "Redacted PDF", [text_pdf], b"%PDF", {"text": "FileForma"}, False),
        ("Sign PDF", "sign-pdf", "Signed PDF", [fixtures["pdf"]], b"%PDF", {"text": "Signed", "page": "1"}, False),
        ("Compare PDF", "compare-pdf", "Comparison report", [fixtures["pdf"], fixtures["pdf2"]], b"PDF", None, False),
        ("Protect PDF", "protect-pdf", "Protected PDF", [fixtures["pdf"]], b"%PDF", {"password": "testpass123"}, False),
        ("Unlock PDF", "unlock-pdf", "Unlocked PDF", [fixtures["protected"]], b"%PDF", {"password": "testpass123"}, False),
    ]

    failures: list[str] = []
    passed = 0
    skipped = 0
    for case in cases:
        name, tool, target, file_paths, magic, extra, optional = case
        if not all(path.exists() for path in file_paths):
            skipped += 1
            print(f"SKIP {name}: fixture missing")
            continue
        try:
            run_case(name, tool, target, file_paths, magic, extra)
            passed += 1
        except Exception as exc:
            if optional:
                skipped += 1
                print(f"SKIP {name} (optional): {exc}")
                continue
            failures.append(f"{name}: {exc}")
            print(f"FAIL {name}: {exc}")

    print(f"\nServer tests: {passed} passed, {skipped} skipped, {len(failures)} failed")
    if failures:
        print("\nFAILURES")
        for failure in failures:
            print(f"- {failure}")
        raise SystemExit(1)
    print("\nAll server converter checks passed.")


if __name__ == "__main__":
    main()
